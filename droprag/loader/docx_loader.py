"""DropRAG DOCX 加载器 - python-docx"""

import os
from typing import Optional, List
from droprag.loader import LoaderBase, LoadedDocument, _get_folder_info, _get_file_mtime
from droprag.logging import get_logger

log = get_logger(__name__)


class DocxLoader(LoaderBase):
    extensions = [".docx", ".doc"]

    def load(self, filepath: str, base_path: str) -> Optional[LoadedDocument]:
        try:
            from docx import Document
        except ImportError:
            log.warning("python-docx 未安装，跳过 DOCX 文件: pip install droprag[office]")
            return None

        try:
            doc = Document(filepath)
        except Exception as e:
            log.debug(f"DOCX 加载失败: {filepath} ({e})")
            return None

        # 提取段落文本
        paragraphs = []
        heading = ""
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # 检测标题
            style_name = para.style.name.lower() if para.style else ""
            if "heading" in style_name and not heading:
                heading = text
            paragraphs.append(text)

        # 提取表格
        tables = []
        for table in doc.tables:
            table_text = self._table_to_text(table)
            if table_text.strip():
                tables.append(table_text)

        content = "\n\n".join(paragraphs)
        if tables:
            content += "\n\n--- 表格 ---\n\n" + "\n\n".join(tables)

        folder, subfolder = _get_folder_info(filepath, base_path)
        return LoadedDocument(
            content=content,
            source=filepath,
            filename=os.path.basename(filepath),
            file_type="docx",
            category="",
            folder=folder,
            subfolder=subfolder,
            file_size=os.path.getsize(filepath),
            file_mtime=_get_file_mtime(filepath),
            heading=heading,
            tables=tables if tables else None,
        )

    @staticmethod
    def _table_to_text(table) -> str:
        """将表格转换为文本格式"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)
